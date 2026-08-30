from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TEMPLATE = DOCS / "课程设计报告模板.docx"
OUTPUT = DOCS / "火车订票系统课程设计报告.docx"
ASSETS = DOCS / "assets"
SCREENSHOTS = DOCS / "rendered" / "screenshots"


def set_font(run, name: str = "宋体", size: float = 12, bold: bool = False, color: str | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 10.5):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.clear()
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def remove_element(element):
    element.getparent().remove(element)


def remove_template_tail(doc):
    heading = next((p for p in doc.paragraphs if "人员组成及分工" in p.text), None)
    if heading is None:
        raise RuntimeError("模板中未找到‘人员组成及分工’标题。")
    body = doc._element.body
    start_removing = False
    for child in list(body):
        if child is heading._element:
            start_removing = True
            continue
        if start_removing and child.tag != qn("w:sectPr"):
            body.remove(child)
    return heading


def replace_cover_slots(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "课程名称" in run.text:
                run.text = run.text.replace("课程名称", "数据库系统")
            if "论文题目" in run.text:
                run.text = run.text.replace("论文题目", "火车订票系统的设计与实现")
    # Some legacy .doc conversions split or retain the placeholders in unusual runs.
    for paragraph in doc.paragraphs:
        if "课程名称" in paragraph.text and "课程论文" in paragraph.text:
            for run in paragraph.runs:
                if "课程名称" in run.text:
                    run.text = run.text.replace("课程名称", "数据库系统")
        if paragraph.text.strip() == "论文题目":
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("火车订票系统的设计与实现")
            set_font(run, name="黑体", size=22, bold=True)


def add_body_paragraph(doc, text: str = "", *, bold: bool = False, before: float = 0, after: float = 6, indent: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, size=12, bold=bold)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, name="宋体", size=14 if level == 1 else 12, bold=True)
    return p


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p


def add_table(doc, headers: list[str], rows: Iterable[Iterable[str]], widths: list[float] | None = None):
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "E9EFF4")
        if widths:
            set_cell_width(cell, widths[index])
    header_props = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_props.append(repeat_header)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER if index in (0, len(cells) - 1) else WD_ALIGN_PARAGRAPH.LEFT)
            if widths:
                set_cell_width(cells[index], widths[index])
    for row in table.rows:
        row.height_rule = None
    return table


def add_figure(doc, image_path: Path, caption: str, width=6.35):
    if not image_path.exists():
        raise RuntimeError(f"缺少报告图片：{image_path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def font(size: int, bold: bool = False):
    for candidate in ("C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf"):
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size, index=0)
    return ImageFont.load_default()


def draw_er_diagram(output: Path):
    output.parent.mkdir(parents=True, exist_ok=True)
    canvas = Image.new("RGB", (2400, 1500), "white")
    draw = ImageDraw.Draw(canvas)
    title_font, box_font, line_font = font(42, True), font(25), font(22, True)
    draw.text((930, 38), "火车订票系统 E-R 图", font=title_font, fill="#1f2933")

    boxes = {
        "用户": (80, 200, 420, 420, ["用户ID（PK）", "用户名", "姓名", "角色"]),
        "订单": (670, 205, 1030, 490, ["订单ID（PK）", "订单号", "用户ID（FK）", "班次ID（FK）", "出发站ID（FK）", "到达站ID（FK）", "状态"]),
        "订单乘车人": (1320, 205, 1720, 455, ["订单乘车人ID（PK）", "订单ID（FK）", "班次座位ID（FK）", "乘车人姓名", "票价"]),
        "班次座位": (1910, 205, 2300, 465, ["班次座位ID（PK）", "班次ID（FK）", "座位ID（FK）", "票价", "状态"]),
        "班次": (690, 760, 1020, 980, ["班次ID（PK）", "车次ID（FK）", "出发日期", "开售状态"]),
        "列车": (120, 1040, 430, 1245, ["列车ID（PK）", "车次编号", "车次名称"]),
        "车厢": (1520, 1040, 1850, 1260, ["车厢ID（PK）", "列车ID（FK）", "车厢号", "席别"]),
        "座位": (1950, 1040, 2260, 1215, ["座位ID（PK）", "车厢ID（FK）", "座位号"]),
        "途经站": (560, 1150, 950, 1410, ["列车ID、车站ID（联合PK）", "站序", "到达时间", "发车时间", "跨天偏移"]),
        "车站": (1030, 1150, 1420, 1370, ["车站ID（PK）", "站点编号", "站点名称", "城市"]),
    }
    def center(rect):
        return ((rect[0] + rect[2]) // 2, (rect[1] + rect[3]) // 2)
    def connector(start, end, left="1", right="N"):
        draw.line((start, end), fill="#7b8794", width=4)
        draw.ellipse((start[0] - 6, start[1] - 6, start[0] + 6, start[1] + 6), fill="#a32112")
        draw.ellipse((end[0] - 6, end[1] - 6, end[0] + 6, end[1] + 6), fill="#1f7a46")
        draw.text((start[0] + 10, start[1] - 32), left, font=line_font, fill="#a32112")
        draw.text((end[0] - 30, end[1] - 32), right, font=line_font, fill="#1f7a46")
    for label, (x1, y1, x2, y2, fields) in boxes.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, outline="#7b8794", width=4, fill="#f8fafb")
        draw.rounded_rectangle((x1, y1, x2, y1 + 52), radius=16, fill="#a32112")
        draw.rectangle((x1, y1 + 32, x2, y1 + 52), fill="#a32112")
        draw.text((x1 + 18, y1 + 10), label, font=box_font, fill="white")
        for index, field in enumerate(fields):
            draw.text((x1 + 20, y1 + 70 + index * 31), field, font=box_font, fill="#24303b")
    connector((420, 310), (670, 310), "1", "N")        # 用户 1—N 订单
    connector((1030, 340), (1320, 340), "1", "N")      # 订单 1—N 订单乘车人
    connector((1720, 330), (1910, 330), "1", "1")      # 订单乘车人 1—1 班次座位
    connector((850, 490), (850, 760), "N", "1")        # 订单 N—1 班次
    connector((1020, 820), (1980, 465), "1", "N")      # 班次 1—N 班次座位
    connector((2105, 1040), (2230, 465), "1", "N")     # 座位 1—N 班次座位
    connector((690, 940), (430, 1150), "N", "1")       # 班次 N—1 列车
    connector((430, 1100), (1520, 1100), "1", "N")     # 列车 1—N 车厢
    connector((1850, 1100), (1950, 1100), "1", "N")    # 车厢 1—N 座位
    connector((430, 1180), (560, 1270), "1", "N")      # 列车 1—N 途经站
    connector((1030, 1300), (950, 1300), "1", "N")     # 车站 1—N 途经站
    canvas.save(output)


def build_report():
    ASSETS.mkdir(parents=True, exist_ok=True)
    er_image = ASSETS / "er-diagram.png"
    draw_er_diagram(er_image)

    doc = Document(TEMPLATE)
    replace_cover_slots(doc)
    remove_template_tail(doc)

    # Personnel page, retained from the template but filled with the agreed equal allocation.
    add_body_paragraph(doc, "本项目由三名成员协作完成。三人按模块交叉评审、共同联调，工作量和贡献度总体均衡；贡献度按 33.33%、33.33%、33.34% 计，合计 100%。", after=10)
    add_table(
        doc,
        ["成员", "主要工作内容", "工作量", "贡献度"],
        [
            ["杨丝宇", "需求分析、乘客端查询与订票界面、登录注册与订单中心页面；负责前后端联调、乘客端测试。", "约 1/3", "33.33%"],
            ["肖永荣", "E-R 图与关系模型设计；MySQL 表、主外键及检查约束；视图、存储过程、触发器和测试数据。", "约 1/3", "33.33%"],
            ["陈逸凡", "管理员后台、统计页面、车站与班次维护；部署说明、测试用例整理、报告排版与答辩材料汇总。", "约 1/3", "33.34%"],
        ],
        [2.1, 9.1, 2.0, 2.0],
    )
    add_body_paragraph(doc, "说明：上述分工以功能模块为主，三人均参与需求评审、代码走查、集成测试和报告校对，因此项目成果由小组共同完成。", before=8, after=0)
    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run("火车订票系统")
    set_font(run, name="宋体", size=16, bold=True)

    add_heading(doc, "一、需求说明")
    add_body_paragraph(doc, "随着铁路网络覆盖范围不断扩大，旅客需要一个能够快速查询车次、比较席别、完成订票并追踪订单状态的应用系统。本项目设计并实现“行程站”火车订票系统，以真实订票业务为背景，围绕用户、车站、列车、班次、座位和订单等数据建立关系模型。系统既满足普通乘客的订票需求，也为管理员提供基础数据维护和销售统计能力。")
    add_heading(doc, "（一）用户角色与业务目标", 2)
    add_table(doc, ["角色", "核心目标", "典型操作"], [
        ["乘客", "快速获取可售车次并完成可靠订票", "注册登录、查询区间车次、选择席别、提交订单、取消订单"],
        ["管理员", "维护基础数据并掌握运营情况", "新增/停用/删除车站、维护班次开售状态、查看订单和日销售统计"],
    ], [2.1, 5.2, 7.9])
    add_heading(doc, "（二）功能需求", 2)
    add_table(doc, ["编号", "功能模块", "功能说明"], [
        ["F01", "身份认证", "乘客可注册、登录和退出；管理员以角色权限进入后台。密码以哈希形式保存。"],
        ["F02", "车次查询", "按出发站、到达站和日期查询班次，展示出发到达时间、余票和最低票价。"],
        ["F03", "在线订票", "乘客选择席别并填写乘车人信息；系统自动分配可售座位，生成唯一订单号。"],
        ["F04", "订单管理", "乘客查看个人订单及座位信息，可取消未取消订单；取消后余票自动恢复。"],
        ["F05", "运营管理", "管理员维护车站和班次开售状态，查看注册乘客、订单、余票和销售额。"],
        ["F06", "统计分析", "通过日销售视图统计已确认订单数、售出票数和销售额，为运营展示提供数据依据。"],
    ], [1.5, 3.0, 10.7])
    add_heading(doc, "（三）非功能需求", 2)
    add_body_paragraph(doc, "系统采用浏览器访问方式，界面需要在桌面和窄屏设备上保持清晰；表单提供必填和错误提示，颜色不作为唯一状态表达。订票操作必须保证同一座位不会被重复出售，因此订单创建过程需要使用数据库事务和行级锁。数据库应具有可维护性，所有关键业务状态均通过枚举、检查约束或触发器进行约束。")

    add_heading(doc, "二、系统设计")
    add_heading(doc, "（一）运行环境与技术选型", 2)
    add_table(doc, ["层次", "采用技术", "作用"], [
        ["前端", "HTML5、CSS3、JavaScript、Jinja 模板", "实现响应式查询、订票、订单和后台数据管理界面。"],
        ["后端", "Python 3.12、Flask 3.1", "提供路由、会话、角色鉴权、输入校验和数据库调用。"],
        ["数据库", "MySQL 8.0、InnoDB、PyMySQL", "存储核心数据，提供事务、外键、视图、存储过程和触发器。"],
        ["开发工具", "VS Code / PowerShell / Microsoft Word", "完成编码、数据库导入、测试、报告排版与文档导出。"],
    ], [2.3, 4.0, 8.9])
    add_heading(doc, "（二）系统架构", 2)
    add_body_paragraph(doc, "系统采用浏览器—Web 应用—数据库三层结构。浏览器通过 Flask 路由访问模板页面；后端完成登录会话、权限判断和业务参数验证；数据层由 MySQL 统一管理。对于查询操作，系统直接读取查询视图；对于订票和取消操作，后端调用存储过程，使关键状态变更集中在数据库事务中执行。")
    add_table(doc, ["模块", "输入", "处理", "输出"], [
        ["车次查询", "出发站、到达站、日期", "读取 v_train_schedule_search 视图", "符合区间条件的可售车次列表"],
        ["订单创建", "班次、区间、席别、乘车人", "调用 sp_create_order；锁定座位并生成订单", "订单号、座位和票价"],
        ["订单取消", "订单号、当前用户", "调用 sp_cancel_order；触发器释放座位", "已取消状态和恢复后的余票"],
        ["后台统计", "已确认订单", "读取 v_daily_sales 视图并聚合", "订单数、票数和销售额"],
    ], [2.3, 4.0, 5.3, 3.6])
    add_heading(doc, "（三）模块说明", 2)
    add_body_paragraph(doc, "乘客端由首页查询、结果列表、席别选择、登录注册和订单中心组成。首页提供站点和日期条件；结果页按照出发时间展示可售车次；席别页显示各席别剩余数量和票价；订单中心提供订单状态和取消操作。管理员后台以运营概览为入口，提供车站维护、车次开售控制、最近订单和日销售统计。")

    add_heading(doc, "三、数据库设计")
    add_heading(doc, "（一）概念结构设计", 2)
    add_body_paragraph(doc, "概念模型以“订单”为核心：一个用户可以拥有多个订单；一个订单包含一个乘车人记录并占用一个班次座位；一个班次属于一趟列车；列车通过途经站实体连接车站，途经站保存站序和到发时刻；列车由车厢和座位构成。班次座位将静态座位映射到具体出发日期，并保存票价和可售状态，从而支持余票查询和并发订票。")
    add_figure(doc, er_image, "图 1 火车订票系统 E-R 图", width=6.35)
    add_heading(doc, "（二）逻辑结构设计", 2)
    add_body_paragraph(doc, "关系模型共包含 10 张表。所有表均使用 InnoDB 引擎；主键保证实体完整性，外键维护参照完整性，检查约束与状态字段维护用户定义完整性。以下表格列出每个关系的主要属性、数据类型和约束说明。")

    table_specs = [
        ("1. 用户表 app_users", [("user_id", "BIGINT", "主键，自增"), ("username", "VARCHAR(32)", "唯一，长度 3—32"), ("password_hash", "VARCHAR(255)", "非空，保存哈希值"), ("real_name", "VARCHAR(32)", "非空"), ("role", "ENUM", "PASSENGER 或 ADMIN")]),
        ("2. 车站表 stations", [("station_id", "BIGINT", "主键，自增"), ("station_code", "VARCHAR(12)", "唯一"), ("station_name", "VARCHAR(50)", "与城市联合唯一"), ("city_name", "VARCHAR(50)", "非空"), ("station_status", "ENUM", "ACTIVE 或 DISABLED")]),
        ("3. 列车表 trains", [("train_id", "BIGINT", "主键，自增"), ("train_code", "VARCHAR(12)", "唯一"), ("train_name", "VARCHAR(80)", "非空"), ("train_status", "ENUM", "ACTIVE 或 DISABLED")]),
        ("4. 列车停靠表 train_stops", [("train_id", "BIGINT", "联合主键、外键"), ("station_id", "BIGINT", "联合主键、外键"), ("stop_sequence", "SMALLINT", "同列车内唯一且大于 0"), ("arrival_time", "TIME", "到站时间"), ("departure_time", "TIME", "发车时间")]),
        ("5. 车厢表 carriages", [("carriage_id", "BIGINT", "主键，自增"), ("train_id", "BIGINT", "外键，关联列车"), ("carriage_no", "SMALLINT", "同列车内唯一"), ("seat_type", "ENUM", "二等座、一等座、商务座")]),
        ("6. 座位表 train_seats", [("seat_id", "BIGINT", "主键，自增"), ("carriage_id", "BIGINT", "外键，关联车厢"), ("seat_no", "VARCHAR(8)", "同车厢内唯一")]),
        ("7. 班次表 train_schedules", [("schedule_id", "BIGINT", "主键，自增"), ("train_id", "BIGINT", "外键，关联列车"), ("travel_date", "DATE", "与列车联合唯一"), ("schedule_status", "ENUM", "ON_SALE、CANCELED、CLOSED")]),
        ("8. 班次座位表 schedule_seats", [("schedule_seat_id", "BIGINT", "主键，自增"), ("schedule_id", "BIGINT", "外键，关联班次"), ("seat_id", "BIGINT", "外键，联合唯一"), ("fare", "DECIMAL(10,2)", "大于 0"), ("seat_status", "ENUM", "AVAILABLE、SOLD、LOCKED")]),
        ("9. 订单表 orders", [("order_id", "BIGINT", "主键，自增"), ("order_no", "VARCHAR(32)", "唯一，订单编号"), ("user_id", "BIGINT", "外键，关联用户"), ("schedule_id", "BIGINT", "外键，关联班次"), ("departure_station_id / arrival_station_id", "BIGINT", "外键，出发/到达站，检查约束保证二者不等"), ("order_status", "ENUM", "PENDING、CONFIRMED、CANCELED"), ("total_amount", "DECIMAL(10,2)", "大于 0，订单总金额")]),
        ("10. 订单乘车人表 order_passengers", [("order_passenger_id", "BIGINT", "主键，自增"), ("order_id", "BIGINT", "外键，关联订单"), ("schedule_seat_id", "BIGINT", "外键，唯一，避免座位重复出票"), ("passenger_name", "VARCHAR(32)", "非空，乘车人姓名"), ("passenger_id_number", "VARCHAR(18)", "非空，乘车人证件号"), ("ticket_price", "DECIMAL(10,2)", "大于 0，成交票价快照")]),
    ]
    for title_text, rows in table_specs:
        add_heading(doc, title_text, 2)
        add_table(doc, ["属性名", "数据类型", "约束说明"], rows, [4.2, 4.1, 8.0])

    add_heading(doc, "（三）规范化程度分析", 2)
    add_body_paragraph(doc, "数据库规范化以函数依赖为判定依据。各表的属性均不可再分，满足第一范式；除列车停靠表使用（列车ID，车站ID）联合主键外，其余表均为单属性代理主键，天然不存在非主属性对候选键的部分函数依赖，满足第二范式。在联合主键表中，站序、到达时间、发车时间和跨天偏移均完全函数依赖于整个候选键，同样不存在部分依赖。")
    add_body_paragraph(doc, "在第三范式检查中，各表的非主属性之间不存在传递函数依赖：席别由车厢直接决定并保存在车厢表中，座位表只保存座位号，未出现“座位→车厢→席别”的冗余链；订单金额直接依赖订单号，乘车人票价保存在乘车人记录中，订单表不重复存储逐票价格；车站、列车的描述属性均直接依赖各自主键。据此，全部 10 张表均满足第三范式，且所有决定因素都是候选键，符合 BCNF 的判定条件。")
    add_body_paragraph(doc, "设计中保留了两处有业务含义的受控冗余：班次座位表的票价和订单乘车人表的票价都是成交当时的快照。铁路票价会随班次和日期调整，若仅通过联表实时计算，历史订单金额将随价格变动而失真；将票价随订单固化保存后，账务可追溯、可对账。两处字段均受检查约束（票价大于 0）保护，并在存储过程事务内一次性写入，保证数据一致。整体设计在满足第三范式的前提下，以少量快照字段换取业务可追溯性，是规范化程度与业务需求之间的合理权衡。")

    add_heading(doc, "（四）数据库高级对象与并发控制", 2)
    add_table(doc, ["对象类型", "名称", "设计作用"], [
        ["视图", "v_train_schedule_search", "将班次、列车、停靠站和可售座位汇总为可按区间查询的车次结果。"],
        ["视图", "v_daily_sales", "以已确认订单为基础统计日销售额、订单数和售出票数。"],
        ["存储过程", "sp_create_order", "在事务中验证区间、使用 FOR UPDATE 锁定可售座位、更新座位状态并生成订单。"],
        ["存储过程", "sp_cancel_order", "验证订单归属与状态后取消订单；后续余票释放交由触发器处理。"],
        ["触发器", "trg_order_passengers_before_insert", "只有座位状态为 SOLD 时才允许写入乘车人记录。"],
        ["触发器", "trg_orders_after_update", "订单变为 CANCELED 时自动将关联班次座位恢复为 AVAILABLE。"],
    ], [2.4, 5.1, 8.8])
    add_body_paragraph(doc, "并发控制方面，订单创建存储过程先开启事务，再使用 SELECT ... FOR UPDATE 对满足条件的最小可售座位执行排他锁。随后以 seat_status='AVAILABLE' 为条件更新座位状态，并检查影响行数是否为 1。若两个用户同时抢购最后一张票，后到的事务会在锁释放后发现没有可售座位并回滚，从而避免超卖。")

    add_heading(doc, "四、系统安装使用说明")
    add_heading(doc, "（一）安装步骤", 2)
    add_table(doc, ["步骤", "操作说明", "预期结果"], [
        ["1", "安装 Python 3.10+ 与 MySQL 8.0，确认 MySQL 服务启动。", "命令行可使用 python 和 mysql。"],
        ["2", "复制 .env.example 为 .env，配置 MYSQL_HOST、MYSQL_USER、MYSQL_PASSWORD 与 FLASK_SECRET_KEY。", "应用能够连接本机 train_ticket_db。"],
        ["3", "执行 scripts/import_database.ps1 -MySqlUser root，依次导入表、存储过程和测试数据。", "生成 train_ticket_db、9 张核心表和演示账号。"],
        ["4", "执行 python -m venv .venv、pip install -r requirements.txt，再运行 python app.py。", "浏览器可打开 http://127.0.0.1:5000。"],
    ], [1.1, 9.7, 5.5])
    add_heading(doc, "（二）主要界面说明", 2)
    add_body_paragraph(doc, "首页将查询条件集中在一个表单中，用户依次选择出发站、到达站和日期；下方展示近期可售行程。车次详情页展示路线、余票、三种席别和乘车人信息。管理员后台集中显示注册乘客、在售车次、已确认订单和销售额，并通过表格呈现日销售统计和最近订单。")
    add_figure(doc, SCREENSHOTS / "home.png", "图 2 乘客端首页与车次查询界面", width=6.35)
    add_figure(doc, SCREENSHOTS / "trip-detail.png", "图 3 席别选择与订票界面", width=6.35)
    add_figure(doc, SCREENSHOTS / "admin-dashboard.png", "图 4 管理员运营概览界面", width=6.35)
    add_heading(doc, "（三）测试用例", 2)
    add_table(doc, ["编号", "测试场景", "操作", "预期结果"], [
        ["T01", "乘客登录", "使用 demo / Demo@123 登录。", "进入首页，导航显示“我的订单”。"],
        ["T02", "车次查询", "查询广州南到北京西的未来日期。", "返回 G1001 及其余票、票价信息。"],
        ["T03", "创建订单", "选择 G1001 二等座并提交乘车人信息。", "生成唯一订单，座位状态由 AVAILABLE 变为 SOLD。"],
        ["T04", "取消订单", "在订单中心取消已确认订单。", "订单状态变为 CANCELED，触发器释放关联座位。"],
        ["T05", "并发抢票", "两个会话同时订购最后一张票。", "最多一个会话成功；另一个收到余票不足提示。"],
        ["T06", "后台统计", "使用 admin / Admin@123 登录后台。", "可查看已确认订单、销售额和 v_daily_sales 统计数据。"],
    ], [1.1, 3.0, 5.4, 6.8])

    add_heading(doc, "五、课程设计总结")
    add_body_paragraph(doc, "本次课程设计以火车订票业务为载体，把数据库设计、Web 编程和界面交互结合起来。项目早期最重要的工作是识别稳定实体和变化实体：列车、车厢和座位属于相对静态数据，而班次座位需要按出发日期单独生成，以承载票价和可售状态。通过这一拆分，系统既能保留列车座位结构，又能准确描述某日某趟车的余票。")
    add_body_paragraph(doc, "在订单处理部分，小组认识到仅靠应用层判断余票无法避免并发超卖。因此将订单创建封装在 MySQL 存储过程中，使用事务、行级锁和状态条件更新确保关键操作具备原子性；订单取消后再由触发器自动恢复余票。该设计使业务规则靠近数据源，也使数据库课程中的视图、存储过程、触发器和完整性约束在系统中得到实际应用。")
    add_body_paragraph(doc, "在系统实现和测试过程中，小组通过角色划分推进工作，又通过代码走查和联调保持接口一致。后续还可以扩展在线支付、多个乘车人、区间分段售票、短信通知和更细粒度的管理员权限。但当前版本已经完成从车次查询到订单、余票、统计的完整闭环，达到了本课程设计的预期目标。")

    add_heading(doc, "参考文献")
    for item in [
        "[1] 王珊，萨师煊. 数据库系统概论（第5版）. 北京：高等教育出版社，2014.",
        "[2] 杨冬青等译. 数据库系统概念（原书第6版）. 北京：机械工业出版社，2012.",
        "[3] 郑阿奇. SQL Server 教程（第3版）. 北京：清华大学出版社，2015.",
        "[4] MySQL 8.0 Reference Manual. Oracle, 2026.",
    ]:
        add_body_paragraph(doc, item, after=3, indent=False)

    add_heading(doc, "附录 A 核心工程文件说明")
    add_table(doc, ["文件或目录", "主要内容", "答辩展示价值"], [
        ["database/schema.sql", "表、主外键、检查约束、视图和触发器定义。", "说明逻辑结构、完整性约束和触发器实现。"],
        ["database/routines.sql", "sp_create_order 与 sp_cancel_order 存储过程。", "演示事务、行级锁和并发控制。"],
        ["database/seed.sql", "车站、车次、座位、班次和演示账号数据。", "快速复现系统初始业务场景。"],
        ["app.py", "Flask 路由、登录会话、角色鉴权和业务调用。", "说明 Web 层如何调用数据库对象。"],
        ["templates / static", "乘客端和管理端页面、响应式样式与交互。", "展示系统界面与操作流程。"],
    ], [4.4, 6.7, 5.2])
    add_body_paragraph(doc, "答辩时可按“管理员登录查看统计 → 乘客查询车次 → 创建订单 → 取消订单并观察余票恢复”的顺序演示系统闭环，再打开数据库视图、存储过程和触发器脚本说明对应实现。", before=8, after=0)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
